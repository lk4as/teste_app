import os
import io
import sys
import tempfile
import streamlit as st
import pandas as pd
from PIL import Image

# -------------------------------
# IMPORTS PARA DOCX
# -------------------------------
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

# Função auxiliar para configurar as bordas de uma célula
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for element in tcPr.xpath('./w:tcBorders'):
        tcPr.remove(element)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_data = kwargs[edge]
            element = OxmlElement(f"w:{edge}")
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
            tcBorders.append(element)
    tcPr.append(tcBorders)

# Configuração padrão de bordas
default_border_settings = {
    "top": {"sz": 4, "val": "single", "color": "000000", "space": "0"},
    "bottom": {"sz": 4, "val": "single", "color": "000000", "space": "0"},
    "left": {"sz": 0, "val": "nil",    "color": "auto",   "space": "0"},
    "right": {"sz": 0, "val": "nil",   "color": "auto",   "space": "0"}
}

def create_chapter_cover(doc, chapter_title):
    if len(doc.paragraphs) > 0:
        doc.add_page_break()
    for _ in range(20):
        p = doc.add_paragraph("")
        p.paragraph_format.line_spacing = 1
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1
    run = para.add_run(chapter_title)
    run.font.size = Pt(20)
    run.font.underline = True
    run.bold = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for _ in range(7):
        p = doc.add_paragraph("")
        p.paragraph_format.line_spacing = 1
    doc.add_page_break()

def create_bordered_section(doc, label, content, no_bottom_border=False,
                            extra_space_top=Pt(3), extra_space_after_content=Pt(3)):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ''
    p_title = cell.add_paragraph()
    p_title.paragraph_format.space_before = extra_space_top
    p_title.paragraph_format.space_after = Pt(0)
    run_title = p_title.add_run(f"{label}:")
    run_title.bold = True
    run_title.font.size = Pt(12)
    p_content = cell.add_paragraph()
    p_content.paragraph_format.space_before = Pt(0)
    p_content.paragraph_format.space_after = extra_space_after_content
    run_content = p_content.add_run(content)
    run_content.font.size = Pt(11)
    if label.lower() == "method":
        p_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        cell.add_paragraph("")
    cell_border_settings = default_border_settings.copy()
    if no_bottom_border:
        cell_border_settings.pop("bottom", None)
    set_cell_border(cell, **cell_border_settings)

def create_test_page(doc, test_info):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    para_title = doc.add_paragraph()
    para_title.paragraph_format.line_spacing = 1
    run_title = para_title.add_run(f"{test_info['Test']}")
    run_title.bold = True
    run_title.font.size = Pt(14)
    para_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("")
    create_bordered_section(doc, "Method", test_info['Method'],
                            extra_space_top=Pt(1), extra_space_after_content=Pt(3))
    doc.add_paragraph("")
    para_steps_title = doc.add_paragraph()
    para_steps_title.paragraph_format.line_spacing = 1
    run_steps_title = para_steps_title.add_run("Steps:")
    run_steps_title.bold = True
    run_steps_title.font.size = Pt(12)
    for step in test_info['Steps']:
        p_step = doc.add_paragraph(f"{step}")
        p_step.paragraph_format.line_spacing = 1
    doc.add_paragraph("")
    create_bordered_section(doc, "Expected Results",
                            "\n".join(test_info['Expected Results']),
                            no_bottom_border=True,
                            extra_space_top=Pt(1),
                            extra_space_after_content=Pt(0))
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    # Max Deviation row
    cells = table.rows[0].cells
    for idx, (label, value_key) in enumerate([
        ("Max Deviation:", "Max. Position Deviation (meters)"),
        ("Heading:",           "Max. Heading Deviation (degrees)")
    ]):
        cell = cells[idx]
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1
        run = p.add_run(label if idx == 0 else "")
        if idx == 0:
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if idx > 0:
            # For the heading column
            run = p.add_run(f" < {test_info[value_key]} degrees")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, **default_border_settings)
    # Results section
    doc.add_paragraph("")
    para_res_title = doc.add_paragraph()
    para_res_title.paragraph_format.line_spacing = 1
    run_res_title = para_res_title.add_run("Results:")
    run_res_title.bold = True
    run_res_title.font.size = Pt(12)
    result_comments = test_info.get('Result + Comment')
    if result_comments and any(pd.notna(r) and str(r).strip() for r in result_comments):
        results_paragraph = doc.add_paragraph()
        results_paragraph.paragraph_format.line_spacing = 1
        for res in result_comments:
            if not pd.notna(res):
                continue
            texto = str(res).strip()
            if texto.lower() == "nan" or texto == "":
                continue
            run_item = results_paragraph.add_run(texto + "\n")
            if "not as expected" in texto.lower():
                run_item.bold = True
    else:
        doc.add_paragraph("")
    # Witness & Date table
    table_info = doc.add_table(rows=2, cols=3)
    hdr_cells = table_info.rows[0].cells
    for idx, hdr in enumerate(["Witness", "Witness", "Date"]):
        p = hdr_cells[idx].paragraphs[0]
        p.paragraph_format.line_spacing = 1
        run = p.add_run(hdr)
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    val_cells = table_info.rows[1].cells
    val_cells[0].paragraphs[0].add_run(str(test_info['Witness 1']))
    val_cells[1].paragraphs[0].add_run(str(test_info['Witness 2']))
    val_cells[2].paragraphs[0].add_run(str(test_info['Date:']))
    for row in table_info.rows:
        for cell in row.cells:
            set_cell_border(cell, **default_border_settings)

def generate_test_report_docx(excel_path):
    df = pd.read_excel(excel_path, sheet_name=0, header=0)
    grouped_tests = {}
    for _, row in df.iterrows():
        num = row['test number']
        section = row['Section']
        if num not in grouped_tests:
            grouped_tests[num] = {
                'Test': row['Test'],
                'Method': row['Method'],
                'Steps': [],
                'Expected Results': [],
                'Result + Comment': [],
                'Max. Position Deviation (meters)': row['Max. Position Deviation (meters)'],
                'Max. Heading Deviation (degrees)': row['Max. Heading Deviation (degrees)'],
                'Witness 1': row['Witness 1'],
                'Witness 2': row['Witness 2'],
                'Date:': row['Date:'],
                'Section': section
            }
        grouped_tests[num]['Steps'].append(row['Step'])
        grouped_tests[num]['Expected Results'].append(row['Expected Result'])
        grouped_tests[num]['Result + Comment'].append(row['Result + Comment'])

    doc = Document()
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Raleway'
    pf = normal_style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(0.2)
        sec.right_margin = Inches(0.2)

    current_chapter = None
    for info in grouped_tests.values():
        if current_chapter != info['Section']:
            create_chapter_cover(doc, info['Section'])
            current_chapter = info['Section']
        create_test_page(doc, info)

    out = os.path.join(tempfile.gettempdir(), "test_report.docx")
    doc.save(out)
    return out

def convert_docx_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path

# -------------------------------
# IMPORTS E FUNÇÕES PARA PDF
# -------------------------------
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

pdfmetrics.registerFont(TTFont('Raleway', 'Raleway-Regular.ttf'))

def create_overlay(page_width, page_height, page_number, overlay_params):
    nome_barco, tipo_teste, mes_ano, rodape_esq, rodape_center, rodape_dir = overlay_params
    packet = io.BytesIO()
    c = canvas.Canvas(packet, pagesize=(page_width, page_height))
    margem = 40
    altura_cabecalho = page_height - 30
    c.setFont("Raleway", 12)
    c.drawString(margem, altura_cabecalho, nome_barco)
    c.drawCentredString(page_width / 2, altura_cabecalho, tipo_teste)
    c.drawRightString(page_width - margem, altura_cabecalho, mes_ano)
    c.setFont("Raleway", 10)
    altura_rodape = 20
    c.drawString(margem, altura_rodape, rodape_esq)
    c.drawCentredString(page_width / 2, altura_rodape, f"{rodape_center} {page_number}")
    c.drawRightString(page_width - margem, altura_rodape, rodape_dir)
    c.save()
    packet.seek(0)
    return PdfReader(packet)

def add_header_footer(input_pdf_path, output_pdf_path, overlay_params):
    reader = PdfReader(input_pdf_path)
    writer = PdfWriter()
    for i, page in enumerate(reader.pages, start=1):
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)
        overlay = create_overlay(w, h, i, overlay_params).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)
    with open(output_pdf_path, "wb") as f:
        writer.write(f)
    return output_pdf_path

def merge_pdfs_func(pdf_list, output_pdf_path):
    merger = PdfMerger()
    for pdf in pdf_list:
        merger.append(pdf)
    merger.write(output_pdf_path)
    merger.close()
    return output_pdf_path

def remove_blank_pages(input_pdf_path, output_pdf_path):
    reader = PdfReader(input_pdf_path)
    writer = PdfWriter()
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            writer.add_page(page)
    with open(output_pdf_path, "wb") as f:
        writer.write(f)
    return output_pdf_path

def run_pdf_merge(doc1_path, doc2_path, overlay_params):
    tmp1 = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    remove_blank_pages(doc2_path, tmp1)
    tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    merge_pdfs_func([doc1_path, tmp1], tmp2)
    final = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    add_header_footer(tmp2, final, overlay_params)
    return final

def get_overlay_params_from_inputs(vessel, test_type, year, abreviation):
    rodape_esq = "Bram DP Assurance"
    rodape_center = ""
    rodape_dir = abreviation
    return vessel, test_type, year, rodape_esq, rodape_center, rodape_dir

# -------------------------------
# INTERFACE COM STREAMLIT
# -------------------------------
st.title("Sistema de Geração de Reports para DP Trials")

# Logo
caminho_imagem = "Logo tradicional.png"
try:
    img = Image.open(caminho_imagem)
    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        st.image(img, use_container_width=True)
except Exception as e:
    st.error(f"Erro ao carregar a imagem. Detalhes: {e}")

tab1, tab2 = st.tabs(["Gerar Relatório", "Mesclar PDFs"])

with tab1:
    st.header("1. Gerar Test Report")
    excel_file = st.file_uploader("Upload da planilha Excel", type=["xlsx"], key="excel_gen")
    if excel_file and st.button("Gerar DOCX"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(excel_file.read())
            path = tmp.name
        docx_path = generate_test_report_docx(path)
        st.success("DOCX gerado!")
        with open(docx_path, "rb") as f:
            st.download_button("Baixar DOCX", f, file_name="test_report.docx")

with tab2:
    st.header("2. Mesclar PDFs")
    st.subheader("Informações para Cabeçalho e Rodapé:")
    vessel       = st.text_input("Nome da embarcação (Vessel)")
    test_type    = st.text_input("Tipo de teste (Type)")
    year         = st.text_input("Ano do teste (Year)")
    abreviation  = st.text_input("Abreviação para rodapé direito (Abreviation)")

    doc1 = st.file_uploader("Upload do PDF da primeira metade (doc1.pdf)", type=["pdf"], key="pdf1")
    doc2 = st.file_uploader("Upload do PDF do Test Report      (doc2.pdf)", type=["pdf"], key="pdf2")

    if st.button("Mesclar PDFs"):
        if doc1 and doc2 and vessel and test_type and year and abreviation:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp1:
                tmp1.write(doc1.read())
                path1 = tmp1.name
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp2:
                tmp2.write(doc2.read())
                path2 = tmp2.name
            overlay = get_overlay_params_from_inputs(vessel, test_type, year, abreviation)
            final_pdf = run_pdf_merge(path1, path2, overlay)
            st.success("PDF final mesclado!")
            with open(final_pdf, "rb") as f:
                st.download_button("Baixar PDF Final", f, file_name="final_report.pdf")
        else:
            st.error("Preencha todos os campos e faça o upload dos dois PDFs.")
